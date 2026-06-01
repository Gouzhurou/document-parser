"""
Парсер для документов Word (.docx, .doc).
"""

import os
import tempfile
from typing import BinaryIO, Union, Optional, List
from pathlib import Path
import pandas as pd

from docx import Document
from docx.table import Table as DocxTable

from ..core.base_parser import BaseParser
from ..core.models import ParseResult, ParsedTable
from ..utils.helpers import (
    normalize_string,
    find_header_rows,
    align_rows_to_max_columns,
    convert_strings_to_dicts,
    clean_empty_rows,
    build_multi_level_headers,
    convert_doc_to_docx_win32,
)
from ..utils.metadata_patterns import extract_metadata_from_text


class WordParser(BaseParser):
    """
    Парсер для документов Word.

    Поддерживает:
    - .docx (напрямую через python-docx)
    - .doc (через конвертацию в .docx)
    """

    def __init__(
            self,
            **kwargs
    ):
        """
        Инициализация парсера Word.

        Args:
            **kwargs: Дополнительные параметры.
        """
        super().__init__(**kwargs)

    def parse(
            self,
            file: Union[BinaryIO, Path, str],
            filename: Optional[str] = None
    ) -> ParseResult:
        """
        Разбирает документ Word и возвращает структурированный результат.
        """
        filename = self._get_filename(file, filename)

        try:
            if isinstance(file, (str, Path)):
                file_path = Path(file)
            else:
                # Сохраняем BinaryIO во временный файл
                suffix = Path(filename).suffix
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    tmp.write(file.read())
                    file_path = Path(tmp.name)

            suffix = file_path.suffix.lower()

            if suffix == '.doc':
                docx_path = convert_doc_to_docx_win32(str(file_path))
                if docx_path:
                    file_path = docx_path

            # Читаем документ
            doc = Document(str(file_path))

            # Извлекаем весь текст для метаданных
            full_text = self._extract_full_text(doc)

            # Извлекаем метаданные
            metadata = extract_metadata_from_text(full_text)

            # Извлекаем таблицы
            tables = self._extract_tables_from_doc(doc)

            # Если таблиц нет, создаём одну с текстом
            if not tables:
                tables = [ParsedTable(
                    sheet_name="Текст документа",
                    headers=["content"],
                    rows=[{"content": line} for line in full_text.split('\n') if line.strip()],
                    metadata=metadata
                )]
            else:
                # Добавляем метаданные к первой таблице
                tables[0].metadata = metadata

            # Удаляем временные файлы
            if isinstance(file, BinaryIO):
                try:
                    os.unlink(file_path)
                except:
                    pass

            return ParseResult(
                filename=filename,
                tables=tables,
                success=True
            )

        except Exception as e:
            import traceback
            traceback.print_exc()

            return ParseResult(
                filename=filename,
                tables=[],
                success=False,
                error=str(e)
            )

    def _extract_full_text(self, doc: Document) -> str:
        """Извлекает весь текст из документа."""
        paragraphs = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        return '\n'.join(paragraphs)

    def _extract_tables_from_doc(self, doc: Document) -> List[ParsedTable]:
        """Извлекает таблицы из документа."""
        tables = []

        for i, table in enumerate(doc.tables):
            parsed = self._parse_docx_table(table, f"Таблица_{i + 1}")
            if parsed and (parsed.headers or parsed.rows):
                tables.append(parsed)

        return tables

    def _parse_docx_table(self, table: DocxTable, sheet_name: str) -> Optional[ParsedTable]:
        """Парсит одну таблицу docx."""
        if not table.rows:
            return None

        # Извлекаем данные из строк
        matrix = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                text = normalize_string(cell.text)
                row_data.append(text)
            matrix.append(row_data)

        matrix = clean_empty_rows(matrix)

        if not matrix:
            return None

        # Конвертируем в DataFrame.
        # Выравниваем строки по максимальному количеству столбцов
        aligned_rows = align_rows_to_max_columns(matrix)

        df = pd.DataFrame(aligned_rows)

        # Находим строки заголовков
        header_start, header_end = find_header_rows(df)

        if header_start is None or header_end is None:
            return None

        # Объединяем строки с заголовками в одну
        headers = build_multi_level_headers(df, (header_start, header_end), len(df.columns))

        data_df = df.iloc[header_end + 1:].copy()

        rows = convert_strings_to_dicts(data_df, headers)

        return ParsedTable(
            sheet_name=sheet_name,
            headers=headers,
            rows=rows
        )
